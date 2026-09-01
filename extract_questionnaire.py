from docx import Document

path = r"D:\论文\Questionnaire.docx"
doc = Document(path)

for i, paragraph in enumerate(doc.paragraphs, 1):
    text = paragraph.text.strip()
    if text:
        print(f"P{i} [{paragraph.style.name}]: {text}")

for ti, table in enumerate(doc.tables, 1):
    print(f"\nTABLE {ti} ({len(table.rows)}x{len(table.columns)})")
    for ri, row in enumerate(table.rows, 1):
        cells = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
        print(f"R{ri}: " + " || ".join(cells))
